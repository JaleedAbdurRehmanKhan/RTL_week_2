from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)

# ── Heading styles ────────────────────────────────────────────────────────────
h1 = doc.styles['Heading 1']
h1.font.size  = Pt(16)
h1.font.bold  = True
h1.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)

h2 = doc.styles['Heading 2']
h2.font.size  = Pt(13)
h2.font.bold  = True
h2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

# ── Helper: add a page break ──────────────────────────────────────────────────
def page_break():
    doc.add_page_break()



# ── Helper: placeholder paragraph ─────────────────────────────────────────────
def placeholder(text="[Screenshot – paste waveform/schematic here]"):
    p = doc.add_paragraph(text)
    p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p.runs[0].font.italic    = True



# ══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("RTL Week 2 Labs")
run.font.size  = Pt(28)
run.font.bold  = True
run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)

doc.add_paragraph()

def centre_field(label, value):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{label}: ")
    r1.font.bold = True
    r1.font.size = Pt(13)
    r2 = p.add_run(value)
    r2.font.size = Pt(13)

centre_field("Date",           "June 2026")
centre_field("Submitted To",   "Sir Khalil")
centre_field("Submitted By",   "Jaleed Abdur Rehman Khan")
centre_field("GitHub Link",    "[INSERT GITHUB LINK HERE]")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  TASKS
# ══════════════════════════════════════════════════════════════════════════════

tasks = [
    {
        "num":   "Task 1",
        "title": "BCD Counter",
        "src":   "bcd_counter.sv",
        "tb":    "tb_bcd_counter.sv",
    },
    {
        "num":   "Task 2",
        "title": "Asynchronous Up/Down Counter",
        "src":   "async_up_down_counter.sv",
        "tb":    "tb_async_up_down_counter.sv",
    },
    {
        "num":   "Task 3",
        "title": "Synchronous Up/Down Counter",
        "src":   "sync_up_down_counter.sv",
        "tb":    "tb_sync_up_down_counter.sv",
    },
    {
        "num":   "Task 4",
        "title": "Mod-N Counter",
        "src":   "mod_n_counter.sv",
        "tb":    "tb_mod_n_counter.sv",
    },
    {
        "num":   "Task 5",
        "title": "Frequency Divider",
        "src":   "freq_div.sv",
        "tb":    "tb_freq_dic.sv",
    },
    {
        "num":   "Task 6",
        "title": "T Flip-Flop",
        "src":   "t_ff.sv",
        "tb":    "tb_t_ff.sv",
    },
    {
        "num":   "Task 7",
        "title": "D Flip-Flop",
        "src":   "d_ff.sv",
        "tb":    "tb_d_ff.sv",
    },
    {
        "num":   "Task 8",
        "title": "JK Flip-Flop",
        "src":   "jk_ff.sv",
        "tb":    "tb_jk_ff.sv",
    },
    {
        "num":   "Task 9",
        "title": "SR Flip-Flop",
        "src":   "sr_ff.sv",
        "tb":    "tb_sr_ff.sv",
    },
    {
        "num":   "Task 10",
        "title": "T Flip-Flop Using D Flip-Flop",
        "src":   "t_using_d_ff.sv",
        "tb":    "tb_t_using_d.sv",
    },
    {
        "num":   "Task 11",
        "title": "SISO – Serial-In Serial-Out Shift Register",
        "src":   "siso.sv",
        "tb":    "tb_siso.sv",
    },
    {
        "num":   "Task 12",
        "title": "SIPO – Serial-In Parallel-Out Shift Register",
        "src":   "sipo.sv",
        "tb":    "tb_sipo.sv",
    },
    {
        "num":   "Task 13",
        "title": "PIPO – Parallel-In Parallel-Out Shift Register",
        "src":   "pipo.sv",
        "tb":    "tb_pipo.sv",
    },
    {
        "num":   "Task 14",
        "title": "PISO – Parallel-In Serial-Out Shift Register",
        "src":   "piso.v",
        "tb":    "tb_piso.sv",
    },
    {
        "num":   "Task 15",
        "title": "Bidirectional Shift Register",
        "src":   "bidirection_shift_reg.sv",
        "tb":    "tb_bidirectonal_shift_reg.sv",
    },
    {
        "num":   "Task 16",
        "title": "Universal Shift Register",
        "src":   "universal_reg.sv",
        "tb":    "tb_universal_shift_reg.sv",
    },
    {
        "num":   "Task 17",
        "title": "Pipelined ALU",
        "src":   "pipelie_alu.sv",
        "tb":    "tb_pipeline_alu.sv",
    },
    {
        "num":   "Task 18",
        "title": "Pipelined Multiplier",
        "src":   "pipeline_multiplier.sv",
        "tb":    "tb_pipeline_multiplier.sv",
    },
]

for i, task in enumerate(tasks):
    # Heading 1
    doc.add_heading(f"{task['num']}: {task['title']}", level=1)

    # Design Code
    doc.add_heading("Design Code", level=2)
    placeholder("[Screenshot – paste design code screenshot here]")

    # Testbench Code
    doc.add_heading("Testbench Code", level=2)
    placeholder("[Screenshot – paste testbench code screenshot here]")

    # Waveform
    doc.add_heading("Waveform", level=2)
    placeholder("[Screenshot – paste simulation waveform here]")

    # Schematic
    doc.add_heading("Schematic", level=2)
    placeholder("[Screenshot – paste elaborated schematic here]")

    # Page break between tasks (not after the last one)
    if i < len(tasks) - 1:
        page_break()

# ── Save ───────────────────────────────────────────────────────────────────────
out_path = r'd:\JaleedUSTP\RTL\RTLLab\week2\RTL_Week2_Lab_Report.docx'
doc.save(out_path)
print(f"Saved: {out_path}")
